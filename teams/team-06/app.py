import streamlit as st
from pypdf import PdfReader
from groq import Groq
from docx import Document
from io import BytesIO
from gtts import gTTS
from datasets import load_dataset
import os
from datetime import datetime
import re

os.makedirs("generated_audio", exist_ok=True)
os.makedirs("generated_docs", exist_ok=True)

# ==========================================
# CONFIG
# ==========================================

GROQ_API_KEY = "API_KEY"
client = Groq(api_key=GROQ_API_KEY)

# ==========================================
# LOAD DATASET
# ==========================================

try:
    dataset = load_dataset(
        "pszemraj/scientific_lay_summarisation-plos-norm",
        split="train[:50]"
    )
    dataset_loaded = True
except Exception as e:
    dataset_loaded = False
    print(e)

# ==========================================
# PAGE CONFIG
# ==========================================

st.set_page_config(
    page_title="Smart Podcast Generation",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# GLOBAL CSS
# ==========================================

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;0,700;0,900;1,400&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

/* ── Reset & Base ── */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

:root {
    --bg-base:        #0A0C10;
    --bg-panel:       #10141C;
    --bg-card:        #161B26;
    --bg-card-hover:  #1C2333;
    --border:         #232B3A;
    --border-bright:  #2E3A50;
    --accent:         #E8B86D;
    --accent-muted:   #C49A4A;
    --accent-glow:    rgba(232, 184, 109, 0.15);
    --red:            #E8697A;
    --green:          #5ECC8B;
    --blue:           #5BA3E8;
    --text-primary:   #EDF2F7;
    --text-secondary: #8A98AC;
    --text-muted:     #4A5568;
    --font-display:   'Playfair Display', Georgia, serif;
    --font-body:      'DM Sans', sans-serif;
    --font-mono:      'DM Mono', monospace;
}

/* ── App Shell ── */
.stApp {
    background: var(--bg-base) !important;
    font-family: var(--font-body) !important;
    color: var(--text-primary) !important;
}

/* Subtle noise texture overlay */
.stApp::before {
    content: '';
    position: fixed;
    inset: 0;
    background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 200 200' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='noise'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.85' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23noise)' opacity='0.04'/%3E%3C/svg%3E");
    pointer-events: none;
    z-index: 0;
    opacity: 0.6;
}

.block-container {
    max-width: 1320px !important;
    padding: 2rem 2.5rem 4rem !important;
    position: relative;
    z-index: 1;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: var(--bg-panel) !important;
    border-right: 1px solid var(--border) !important;
}

[data-testid="stSidebar"] .block-container {
    padding: 2rem 1.5rem !important;
}

[data-testid="stSidebar"] * {
    color: var(--text-secondary) !important;
}

[data-testid="stSidebarContent"] > div:first-child {
    margin-top: 0 !important;
}

/* ── Typography ── */
h1, h2, h3, .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
    font-family: var(--font-display) !important;
    font-weight: 700 !important;
    letter-spacing: -0.02em !important;
    color: var(--text-primary) !important;
}

p, li, .stMarkdown p, .stMarkdown li {
    font-family: var(--font-body) !important;
    color: var(--text-secondary) !important;
    line-height: 1.85 !important;
}

/* ── Hero Header ── */
.podpaper-hero {
    background: linear-gradient(135deg, #0D111A 0%, #141A28 60%, #0F131C 100%);
    border: 1px solid var(--border-bright);
    border-radius: 20px;
    padding: 3.5rem 4rem;
    margin-bottom: 2.5rem;
    position: relative;
    overflow: hidden;
}

.podpaper-hero::before {
    content: '';
    position: absolute;
    top: -80px; right: -80px;
    width: 320px; height: 320px;
    background: radial-gradient(circle, rgba(232,184,109,0.12) 0%, transparent 70%);
    border-radius: 50%;
    pointer-events: none;
}

.podpaper-hero::after {
    content: '';
    position: absolute;
    bottom: -60px; left: -60px;
    width: 240px; height: 240px;
    background: radial-gradient(circle, rgba(91,163,232,0.07) 0%, transparent 70%);
    border-radius: 50%;
    pointer-events: none;
}

.hero-eyebrow {
    font-family: var(--font-mono) !important;
    font-size: 0.7rem !important;
    letter-spacing: 0.2em !important;
    text-transform: uppercase !important;
    color: var(--accent) !important;
    margin-bottom: 1rem !important;
    display: flex;
    align-items: center;
    gap: 8px;
}

.hero-eyebrow::before {
    content: '';
    display: inline-block;
    width: 24px; height: 1px;
    background: var(--accent);
}

.hero-title {
    font-family: var(--font-display) !important;
    font-size: clamp(2.4rem, 5vw, 4rem) !important;
    font-weight: 900 !important;
    line-height: 1.05 !important;
    letter-spacing: -0.04em !important;
    color: var(--text-primary) !important;
    margin-bottom: 1.2rem !important;
}

.hero-title span {
    color: var(--accent) !important;
    font-style: italic !important;
}

.hero-subtitle {
    font-family: var(--font-body) !important;
    font-size: 1.05rem !important;
    color: var(--text-secondary) !important;
    line-height: 1.7 !important;
    max-width: 560px !important;
    font-weight: 300 !important;
}

/* ── Section Labels ── */
.section-label {
    font-family: var(--font-mono) !important;
    font-size: 0.65rem !important;
    letter-spacing: 0.18em !important;
    text-transform: uppercase !important;
    color: var(--accent-muted) !important;
    margin-bottom: 1rem !important;
    display: flex;
    align-items: center;
    gap: 10px;
}

.section-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: var(--border);
}

/* ── Cards ── */
.info-card {
    background: var(--bg-card);
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 1.5rem 1.75rem;
    transition: border-color 0.25s ease, background 0.25s ease;
}

.info-card:hover {
    border-color: var(--border-bright);
    background: var(--bg-card-hover);
}

.stack-row {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 0.55rem 0;
    border-bottom: 1px solid var(--border);
}

.stack-row:last-child { border-bottom: none; }

.stack-icon {
    width: 28px; height: 28px;
    background: var(--bg-base);
    border: 1px solid var(--border-bright);
    border-radius: 6px;
    display: flex; align-items: center; justify-content: center;
    font-size: 0.75rem;
    flex-shrink: 0;
}

.stack-label {
    font-family: var(--font-mono) !important;
    font-size: 0.78rem !important;
    color: var(--text-secondary) !important;
}

.stack-value {
    font-family: var(--font-body) !important;
    font-size: 0.78rem !important;
    color: var(--text-muted) !important;
    margin-left: auto !important;
}

/* ── Status Badge ── */
.status-badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: rgba(94, 204, 139, 0.1);
    border: 1px solid rgba(94, 204, 139, 0.25);
    color: var(--green) !important;
    font-family: var(--font-mono) !important;
    font-size: 0.7rem !important;
    letter-spacing: 0.08em !important;
    padding: 4px 12px !important;
    border-radius: 100px !important;
    margin-top: 1rem !important;
}

.status-badge::before {
    content: '●';
    font-size: 0.5rem;
    animation: pulse 2s infinite;
}

@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.3; }
}

/* ── Selectbox ── */
.stSelectbox > div > div {
    background: var(--bg-card) !important;
    border: 1px solid var(--border-bright) !important;
    border-radius: 10px !important;
    color: var(--text-primary) !important;
    font-family: var(--font-body) !important;
    font-size: 0.9rem !important;
}

.stSelectbox > div > div:hover {
    border-color: var(--accent) !important;
}

.stSelectbox label, .stFileUploader label {
    font-family: var(--font-mono) !important;
    font-size: 0.68rem !important;
    letter-spacing: 0.14em !important;
    text-transform: uppercase !important;
    color: var(--text-muted) !important;
}

/* Dropdown options */
[data-baseweb="popover"] {
    background: var(--bg-panel) !important;
    border: 1px solid var(--border-bright) !important;
    border-radius: 10px !important;
}

[role="option"] {
    color: var(--text-secondary) !important;
    font-family: var(--font-body) !important;
}

[role="option"]:hover {
    background: var(--bg-card-hover) !important;
    color: var(--text-primary) !important;
}

/* ── File Uploader ── */
[data-testid="stFileUploader"] {
    background: var(--bg-card) !important;
    border: 2px dashed var(--border-bright) !important;
    border-radius: 16px !important;
    padding: 2.5rem !important;
    text-align: center !important;
    transition: border-color 0.25s ease, background 0.25s ease !important;
}

[data-testid="stFileUploader"]:hover {
    border-color: var(--accent) !important;
    background: var(--bg-card-hover) !important;
}

[data-testid="stFileUploader"] * {
    color: var(--text-secondary) !important;
}

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, var(--accent) 0%, var(--accent-muted) 100%) !important;
    color: #0A0C10 !important;
    font-family: var(--font-mono) !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.8rem 2.2rem !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 20px rgba(232, 184, 109, 0.25) !important;
    width: 100% !important;
}

.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 30px rgba(232, 184, 109, 0.4) !important;
}

.stButton > button:active {
    transform: translateY(0) !important;
}

/* ── Download Buttons ── */
.stDownloadButton > button {
    background: var(--bg-card) !important;
    color: var(--text-primary) !important;
    border: 1px solid var(--border-bright) !important;
    border-radius: 10px !important;
    font-family: var(--font-mono) !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
    padding: 0.6rem 1.5rem !important;
    transition: all 0.2s ease !important;
    box-shadow: none !important;
    width: auto !important;
}

.stDownloadButton > button:hover {
    background: var(--bg-card-hover) !important;
    border-color: var(--accent) !important;
    color: var(--accent) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(232,184,109,0.15) !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    background: var(--bg-card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text-secondary) !important;
    font-family: var(--font-mono) !important;
    font-size: 0.75rem !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
}

.streamlit-expanderContent {
    background: var(--bg-panel) !important;
    border: 1px solid var(--border) !important;
    border-top: none !important;
    border-radius: 0 0 10px 10px !important;
    padding: 1.5rem !important;
    font-family: var(--font-mono) !important;
    font-size: 0.8rem !important;
    color: var(--text-muted) !important;
    white-space: pre-wrap !important;
    line-height: 1.7 !important;
}

/* ── Metrics ── */
[data-testid="stMetric"] {
    background: var(--bg-card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
    padding: 1.4rem 1.75rem !important;
}

[data-testid="stMetricLabel"] {
    font-family: var(--font-mono) !important;
    font-size: 0.66rem !important;
    letter-spacing: 0.14em !important;
    text-transform: uppercase !important;
    color: var(--text-muted) !important;
}

[data-testid="stMetricValue"] {
    font-family: var(--font-display) !important;
    font-size: 2.4rem !important;
    font-weight: 700 !important;
    color: var(--accent) !important;
}

/* ── Alerts ── */
.stSuccess {
    background: rgba(94,204,139,0.08) !important;
    border: 1px solid rgba(94,204,139,0.25) !important;
    border-radius: 10px !important;
    color: var(--green) !important;
}

.stError {
    background: rgba(232,105,122,0.08) !important;
    border: 1px solid rgba(232,105,122,0.25) !important;
    border-radius: 10px !important;
    color: var(--red) !important;
}

.stWarning {
    background: rgba(232,184,109,0.08) !important;
    border: 1px solid rgba(232,184,109,0.2) !important;
    border-radius: 10px !important;
    color: var(--accent) !important;
}

/* ── Spinner ── */
.stSpinner > div {
    border-top-color: var(--accent) !important;
}

/* ── Audio Player ── */
audio {
    width: 100% !important;
    border-radius: 10px !important;
    accent-color: var(--accent) !important;
    background: var(--bg-card) !important;
}

/* ── Output Sections ── */
.output-block {
    background: var(--bg-card);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 2rem 2.25rem;
    margin: 1.5rem 0;
    position: relative;
    overflow: hidden;
}

.output-block::before {
    content: '';
    position: absolute;
    top: 0; left: 0;
    width: 3px; height: 100%;
    background: linear-gradient(180deg, var(--accent) 0%, transparent 100%);
    border-radius: 3px 0 0 3px;
}

.output-block h3 {
    font-family: var(--font-display) !important;
    font-size: 1.3rem !important;
    color: var(--text-primary) !important;
    margin-bottom: 1rem !important;
}

.output-block p, .output-block li {
    font-family: var(--font-body) !important;
    font-size: 0.9rem !important;
    color: var(--text-secondary) !important;
    line-height: 1.85 !important;
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* ── Subheaders ── */
.stSubheader, [data-testid="stSubheader"] {
    font-family: var(--font-display) !important;
    font-size: 1.5rem !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    border-bottom: 1px solid var(--border) !important;
    padding-bottom: 0.75rem !important;
    margin: 2rem 0 1.25rem !important;
}

/* ── Podcast Script Formatting ── */
.podcast-output {
    font-family: var(--font-body) !important;
    font-size: 0.93rem !important;
    line-height: 2 !important;
    color: var(--text-secondary) !important;
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* ── Divider ── */
hr {
    border: none !important;
    border-top: 1px solid var(--border) !important;
    margin: 2rem 0 !important;
}

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: var(--bg-base); }
::-webkit-scrollbar-thumb { background: var(--border-bright); border-radius: 10px; }
::-webkit-scrollbar-thumb:hover { background: var(--text-muted); }

/* ── Justify ALL Text Content ── */
/* Markdown paragraphs and lists */
.stMarkdown p,
.stMarkdown li,
.stMarkdown ol li,
.stMarkdown ul li,
div[data-testid="stMarkdownContainer"] p,
div[data-testid="stMarkdownContainer"] li {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Output blocks (analysis, gap, podcast) */
.output-block p,
.output-block li,
.output-block div,
.podcast-output,
.podcast-output p,
.podcast-output li {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Expander content (extracted text preview) */
.streamlit-expanderContent p,
.streamlit-expanderContent div,
.streamlit-expanderContent span {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Sidebar text */
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] div,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] .stMarkdown p {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Success / warning / error / info alerts */
.stSuccess, .stError, .stWarning, .stInfo,
[data-testid="stNotification"] {
    text-align: justify !important;
    text-justify: inter-word !important;
}

/* Generic Streamlit text elements */
.stText, [data-testid="stText"],
div.element-container p,
div.element-container span,
div.element-container li {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Write / st.write output */
div[data-testid="stWrite"] p,
div[data-testid="stWrite"] li {
    text-align: justify !important;
    text-justify: inter-word !important;
    hyphens: auto !important;
}

/* Headings stay left-aligned for visual hierarchy */
h1, h2, h3, h4, h5, h6,
.stMarkdown h1, .stMarkdown h2,
.stMarkdown h3, .stMarkdown h4 {
    text-align: left !important;
}

/* ── Columns gap ── */
.row-widget { gap: 1.25rem !important; }

/* ── Markdown code blocks ── */
code {
    font-family: var(--font-mono) !important;
    background: var(--bg-base) !important;
    border: 1px solid var(--border) !important;
    border-radius: 5px !important;
    padding: 0.15em 0.5em !important;
    font-size: 0.8em !important;
    color: var(--accent) !important;
}

pre code {
    padding: 1rem !important;
    display: block !important;
}

</style>
""", unsafe_allow_html=True)


# ==========================================
# HERO SECTION
# ==========================================

st.markdown("""
<div class="podpaper-hero">
    <div class="hero-eyebrow">AI-Powered Research Podcast Studio</div>
    <div class="hero-title">Turn Papers Into<br><span>Podcasts</span></div>
    <div class="hero-subtitle">
        Upload any research paper and watch PodPaper AI transform dense academic prose
        into a compelling, ready-to-publish podcast script — complete with audio, timestamps, and a DOCX export.
    </div>
</div>
""", unsafe_allow_html=True)


# ==========================================
# SIDEBAR
# ==========================================

with st.sidebar:
    st.markdown("""
    <div style="margin-bottom:2rem;">
        <div style="font-family:'Playfair Display',serif; font-size:1.4rem; font-weight:900;
                    color:#EDF2F7; letter-spacing:-0.03em; margin-bottom:0.25rem;">
            ResearchCast<span style="color:#E8B86D;"></span>
        </div>
        <div style="font-family:'DM Mono',monospace; font-size:0.65rem; letter-spacing:0.14em;
                    text-transform:uppercase; color:#4A5568;">
            Research → Podcast
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="section-label">Tech Stack</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-card" style="margin-bottom:1.5rem;">
        <div class="stack-row">
            <div class="stack-icon">🧠</div>
            <span class="stack-label">Model</span>
            <span class="stack-value">Llama 3.3 70B</span>
        </div>
        <div class="stack-row">
            <div class="stack-icon">⚡</div>
            <span class="stack-label">Inference</span>
            <span class="stack-value">Groq</span>
        </div>
        <div class="stack-row">
            <div class="stack-icon">📄</div>
            <span class="stack-label">PDF Parsing</span>
            <span class="stack-value">PyPDF</span>
        </div>
        <div class="stack-row">
            <div class="stack-icon">🖥️</div>
            <span class="stack-label">Frontend</span>
            <span class="stack-value">Streamlit</span>
        </div>
        <div class="stack-row">
            <div class="stack-icon">📦</div>
            <span class="stack-label">Exports</span>
            <span class="stack-value">DOCX · MP3</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="section-label">Dataset</div>', unsafe_allow_html=True)

    if dataset_loaded:
        st.markdown(f"""
        <div class="info-card">
            <div style="font-family:'DM Sans',sans-serif; font-size:0.82rem; color:#8A98AC; line-height:1.7;">
                Scientific Lay Summarisation (PLOS) — {len(dataset)} samples loaded to guide accessible science communication.
            </div>
            <div class="status-badge">Dataset Active</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="info-card">
            <div style="font-family:'DM Sans',sans-serif; font-size:0.82rem; color:#E8697A;">
                Dataset unavailable — running without examples.
            </div>
        </div>
        """, unsafe_allow_html=True)


# ==========================================
# MAIN CONTROLS
# ==========================================

col_left, col_right = st.columns([1, 1.6], gap="large")

with col_left:
    st.markdown('<div class="section-label">Podcast Style</div>', unsafe_allow_html=True)
    style = st.selectbox(
        "Select a style",
        ["Educational", "Casual", "Technical", "Student Friendly", "Research Seminar"],
        label_visibility="collapsed"
    )

    style_descriptions = {
        "Educational":       "Clear, structured, and informative. Great for general audiences.",
        "Casual":            "Relaxed and conversational. Like chatting with a friend.",
        "Technical":         "In-depth and precise. Perfect for domain experts.",
        "Student Friendly":  "Simplified language with real-world analogies.",
        "Research Seminar":  "Academic rigour with a professional seminar tone.",
    }

    st.markdown(f"""
    <div style="margin-top:0.75rem; padding:1rem 1.25rem; background:var(--bg-card);
                border:1px solid var(--border); border-radius:10px;">
        <div style="font-family:'DM Mono',monospace; font-size:0.65rem; letter-spacing:0.1em;
                    text-transform:uppercase; color:#4A5568; margin-bottom:0.4rem;">Style Preview</div>
        <div style="font-family:'DM Sans',sans-serif; font-size:0.83rem; color:#8A98AC; line-height:1.6;">
            {style_descriptions[style]}
        </div>
    </div>
    """, unsafe_allow_html=True)

with col_right:
    st.markdown('<div class="section-label">Upload Paper</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Drop your PDF here",
        type=["pdf"],
        label_visibility="collapsed"
    )


# ==========================================
# PROCESSING
# ==========================================

if uploaded_file is not None:

    try:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text

        # Success banner
        st.markdown("""
        <div style="display:flex; align-items:center; gap:10px; margin:1.5rem 0 0.5rem;
                    padding:0.85rem 1.25rem; background:rgba(94,204,139,0.07);
                    border:1px solid rgba(94,204,139,0.22); border-radius:10px;">
            <span style="color:#5ECC8B; font-size:1.1rem;">✓</span>
            <span style="font-family:'DM Mono',monospace; font-size:0.73rem; letter-spacing:0.1em;
                         text-transform:uppercase; color:#5ECC8B;">
                PDF Processed Successfully
            </span>
            <span style="margin-left:auto; font-family:'DM Mono',monospace; font-size:0.7rem;
                         color:#4A5568;">
                {word_count} words extracted
            </span>
        </div>
        """.format(word_count=len(text.split())), unsafe_allow_html=True)

        with st.expander("▸  Preview Extracted Text"):
            st.write(text[:3000])

        # Generate button
        st.markdown("<div style='margin:1.5rem 0 0.5rem;'>", unsafe_allow_html=True)
        generate = st.button("🎙️  Generate Podcast Script")
        st.markdown("</div>", unsafe_allow_html=True)

        if generate:

            with st.spinner("Synthesising your podcast — this takes about 30 seconds…"):

                # ──────────────────────────────
                # PAPER ANALYSIS
                # ──────────────────────────────

                analysis_prompt = f"""
Analyze this research paper.

Return:

1. Paper Title
2. Research Domain
3. Main Problem
4. Methodology
5. Dataset Used
6. Evaluation Metrics
7. Key Results
8. Future Scope

Research Paper:

{text[:12000]}
"""

                analysis_response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": analysis_prompt}]
                )
                paper_analysis = analysis_response.choices[0].message.content

            st.markdown('<div class="section-label" style="margin-top:2.5rem;">Paper Analysis</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="output-block">{paper_analysis}</div>', unsafe_allow_html=True)

            with st.spinner("Identifying research gaps…"):

                # ──────────────────────────────
                # RESEARCH GAP
                # ──────────────────────────────

                gap_prompt = f"""
Research Paper:

{text[:10000]}

Identify:

1. Research Gap
2. Novel Contribution
3. Difference From Existing Work
4. Future Research Opportunities
"""

                gap_response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": gap_prompt}]
                )
                research_gap = gap_response.choices[0].message.content

            st.markdown('<div class="section-label" style="margin-top:2rem;">Research Gap Analysis</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="output-block">{research_gap}</div>', unsafe_allow_html=True)

            # ──────────────────────────────
            # DATASET SECTION
            # ──────────────────────────────

            dataset_examples = ""
            if dataset_loaded:
                st.markdown('<div class="section-label" style="margin-top:2rem;">Dataset Context</div>', unsafe_allow_html=True)
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    st.metric("Samples Loaded", len(dataset))
                with col_d2:
                    st.metric("Examples Used", 3)

                st.markdown("""
                <div style="margin-top:1rem; padding:1rem 1.25rem; background:var(--bg-card);
                            border:1px solid var(--border); border-radius:10px;
                            font-family:'DM Sans',sans-serif; font-size:0.85rem;
                            color:#8A98AC; line-height:1.75;">
                    The <strong style="color:#EDF2F7;">Scientific Lay Summarisation</strong> dataset contains
                    scientific papers paired with human-written layman summaries.
                    These examples guide the AI in translating complex research into accessible podcast language.
                </div>
                """, unsafe_allow_html=True)

                for row in dataset.select(range(3)):
                    dataset_examples += f"""

Title:
{row['title']}

Article:
{row['article'][:800]}

Lay Summary:
{row['summary'][:300]}

"""

            with st.spinner("Writing your podcast script…"):

                # ──────────────────────────────
                # PODCAST GENERATION
                # ──────────────────────────────

                podcast_prompt = f"""
You are an expert science communicator.

Podcast Style: {style}

Reference examples from Scientific Lay Summarisation Dataset:

{dataset_examples}

Learn how scientific papers are converted into
simple explanations from these examples.

Generate:

# Podcast Title

# Podcast Script

Requirements:

- Host and Research Expert conversation
- Do not invent names
- Use only:

Host:
Research Expert:

- Conversational style
- Easy language
- Explain:
  - Problem
  - Methodology
  - Experiments
  - Results
  - Limitations
  - Future Work

- Include 3 analogies naturally

# Three Listener-Friendly Analogies

# Show Notes

# Key Terms

# Timestamps

Generate podcast chapter timestamps.

Requirements:

- Use MM:SS format
- Use start and end times
- Cover the entire podcast
- Generate 6-10 chapters
- Use realistic durations

Example:

00:00 - 00:45 Welcome
00:45 - 02:30 Research Problem
02:30 - 05:10 Methodology
05:10 - 07:00 Results
07:00 - 08:00 Conclusion

# Key Takeaways

Research Paper:

{text[:25000]}
"""

                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": podcast_prompt}],
                    temperature=0.7,
                    max_tokens=4000
                )
                output = response.choices[0].message.content

            # ──────────────────────────────
            # METRICS ROW
            # ──────────────────────────────

            word_count = len(output.split())
            duration = round(word_count / 150)

            st.markdown('<div class="section-label" style="margin-top:2.5rem;">Podcast Stats</div>', unsafe_allow_html=True)
            mc1, mc2, mc3 = st.columns(3)
            with mc1:
                st.metric("Word Count", f"{word_count:,}")
            with mc2:
                st.metric("Est. Duration", f"{duration} min")
            with mc3:
                st.metric("Style", style)

            # ──────────────────────────────
            # PODCAST SCRIPT
            # ──────────────────────────────

            output = re.sub(r'(\d{2}:\d{2}\s*-\s*\d{2}:\d{2})', r'\n\n\1', output)

            st.markdown('<div class="section-label" style="margin-top:2.5rem;">🎙️ Generated Podcast</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="output-block podcast-output">{output}</div>', unsafe_allow_html=True)

            # ──────────────────────────────
            # EXPORTS
            # ──────────────────────────────

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            st.markdown('<div class="section-label" style="margin-top:2rem;">Export</div>', unsafe_allow_html=True)

            exp1, exp2 = st.columns(2)

            # DOCX
            doc = Document()
            doc.add_heading("Podcast Script", level=1)
            doc.add_paragraph(output)
            doc_path = f"generated_docs/podcast_{timestamp}.docx"
            doc.save(doc_path)
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            with exp1:
                st.download_button(
                    label="📥  Download DOCX",
                    data=doc_buffer,
                    file_name=f"podcast_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # AUDIO
            try:
                audio_text = output[:5000]
                tts = gTTS(text=audio_text, lang="en")
                audio_file = f"generated_audio/podcast_{timestamp}.mp3"
                tts.save(audio_file)

                with open(audio_file, "rb") as audio:
                    audio_bytes = audio.read()

                st.markdown('<div class="section-label" style="margin-top:2rem;">🎧 Podcast Audio</div>', unsafe_allow_html=True)
                st.audio(audio_bytes)

                with exp2:
                    st.download_button(
                        label="⬇  Download MP3",
                        data=audio_bytes,
                        file_name=f"podcast_{timestamp}.mp3",
                        mime="audio/mpeg"
                    )

            except Exception as audio_error:
                st.warning(f"Audio generation failed: {audio_error}")

    except Exception as e:
        st.error(f"PDF Processing Error: {e}")

else:
    # ──────────────────────────────
    # EMPTY STATE
    # ──────────────────────────────

    st.markdown("""
    <div style="margin-top:3rem; text-align:center; padding:4rem 2rem;
                background:var(--bg-card); border:1px dashed var(--border-bright);
                border-radius:20px;">
        <div style="font-size:3rem; margin-bottom:1.25rem; opacity:0.5;">📄</div>
        <div style="font-family:'Playfair Display',serif; font-size:1.5rem;
                    font-weight:700; color:#EDF2F7; margin-bottom:0.6rem;">
            No paper uploaded yet
        </div>
        <div style="font-family:'DM Sans',sans-serif; font-size:0.88rem;
                    color:#4A5568; max-width:380px; margin:0 auto; line-height:1.75;">
            Upload a research paper PDF above, choose a podcast style, and let
            PodPaper AI do the rest.
        </div>
    </div>
    """, unsafe_allow_html=True)