import streamlit as st
import pandas as pd
import PyPDF2
from docx import Document
from langchain_community.llms import Ollama
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
import io

df = pd.read_csv("insurance_claim_rejections_dataset.csv")
llm = Ollama(model="llama3")

prompt_template = PromptTemplate(
    input_variables=["rejection_text", "rejection_reason", "policy_clause", "plain_explanation", "appeal_suggestion"],
    template="""
You are an insurance expert helping a policyholder.
Rejection Letter: {rejection_text}
Known Rejection Reason: {rejection_reason}
Policy Clause: {policy_clause}
Plain English Explanation: {plain_explanation}
Appeal Suggestion: {appeal_suggestion}
1. Explain the rejection in very simple plain language (3-4 lines).
2. Write a formal structured appeal letter.
"""
)

chain = prompt_template | llm | StrOutputParser()

def extract_text_from_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def match_from_dataset(claim_type, rejection_text):
    filtered = df[df['Claim_Type'].str.lower() == claim_type.lower()]
    for _, row in filtered.iterrows():
        if any(word in rejection_text.lower() for word in row['Rejection_Reason'].lower().split()):
            return row
    return filtered.iloc[0]

def save_as_docx(content, claim_type):
    doc = Document()
    doc.add_heading("Insurance Claim Appeal Letter", 0)
    doc.add_paragraph(f"Claim Type: {claim_type}")
    doc.add_paragraph("---")
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.title("🛡️ Insurance Claim Rejection Explainer & Appeal Drafter")
st.markdown("Upload your rejection letter and get a plain language explanation + appeal letter!")

claim_type = st.selectbox("Select Claim Type", ["Health", "Motor", "Life"])
uploaded_file = st.file_uploader("Upload Rejection Letter (PDF)", type=["pdf"])
manual_text = st.text_area("Or paste rejection letter text here")

if st.button("Analyse & Generate Appeal"):
    if uploaded_file:
        rejection_text = extract_text_from_pdf(uploaded_file)
    elif manual_text:
        rejection_text = manual_text
    else:
        st.error("Please upload a PDF or paste text!")
        st.stop()

    with st.spinner("Analysing with AI..."):
        matched = match_from_dataset(claim_type, rejection_text)
        result = chain.invoke({
            "rejection_text": rejection_text,
            "rejection_reason": matched['Rejection_Reason'],
            "policy_clause": matched['Policy_Clause'],
            "plain_explanation": matched['Plain_English_Explanation'],
            "appeal_suggestion": matched['Appeal_Suggestion']
        })

    st.subheader("📋 Policy Clause Referenced")
    st.info(matched['Policy_Clause'])

    st.subheader("💡 AI Explanation & Appeal Letter")
    st.write(result)

    st.subheader("📥 Download Appeal Letter")
    docx_file = save_as_docx(result, claim_type)
    st.download_button("Download as DOCX", docx_file, file_name="appeal_letter.docx")