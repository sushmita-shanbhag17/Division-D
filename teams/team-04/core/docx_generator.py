# DOCX File Generator for Court Submission
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import re

def build_docx_from_markdown(petition_markdown: str) -> io.BytesIO:
    """
    Parses petition markdown and returns a formatted Word Document as a byte stream.
    """
    doc = docx.Document()
    
    # Page Setup (Standard Legal Format: 1-inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Custom Heading 1 Style
    try:
        style_h1 = doc.styles.add_style('LegalHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = style_h1.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(14)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        style_h1.paragraph_format.space_before = Pt(12)
        style_h1.paragraph_format.space_after = Pt(6)
        style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        style_h1 = doc.styles['Heading 1']
        
    # Custom Heading 2 Style
    try:
        style_h2 = doc.styles.add_style('LegalHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = style_h2.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        style_h2.paragraph_format.space_before = Pt(12)
        style_h2.paragraph_format.space_after = Pt(6)
        style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        style_h2 = doc.styles['Heading 2']

    # Parse and write markdown line by line
    lines = petition_markdown.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
            
        clean_text = line
        lower_line = line.lower()
        
        # Determine paragraph style and alignment based on content
        style_name = 'Normal'
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent = None
        
        if line.startswith('# '):
            clean_text = line[2:].strip()
            style_name = 'LegalHeading1'
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            clean_text = line[3:].strip()
            style_name = 'LegalHeading1'
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('### ') or line.startswith('#### '):
            clean_text = re.sub(r'^#+\s+', '', line).strip()
            style_name = 'LegalHeading2'
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif 'date:' in lower_line and len(line) < 40:
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'before the' in lower_line or 'commission' in lower_line or 'forum' in lower_line or 'petition no' in lower_line:
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif lower_line in ['versus', 'vs', 'vs.', '**versus**']:
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif lower_line.startswith('in the matter of') or lower_line.startswith('__in the matter of'):
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif lower_line.strip().startswith('.... complainant') or lower_line.strip().startswith('.... petitioner') or lower_line.strip().startswith('.......respondent'):
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'subject:' in lower_line or 'sub:' in lower_line:
            indent = Inches(0.5)
            if not line.startswith('**') and not line.startswith('__'):
                clean_text = f"**{clean_text}**"
        elif line.startswith('- ') or line.startswith('* '):
            clean_text = line[2:].strip()
            indent = Inches(0.5)
        elif re.match(r'^[\*_]*\d+(?:\.\d+)*\.?\s', line):
            indent = Inches(0.25)
        elif len(line) < 80 and line.isupper() and not line.startswith('**') and not line.startswith('__'):
            alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean_text = f"__{clean_text}__"
            
        # Create paragraph
        p = doc.add_paragraph(style=style_name)
        p.alignment = alignment
        if indent:
            p.paragraph_format.left_indent = indent
            
        _add_formatted_text(p, clean_text)
        i += 1
        
    # Save to BytesIO
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

def _add_formatted_text(paragraph, text):
    """
    Parses basic bold formatting (**text**) and underline (__text__) and adds to paragraph runs.
    """
    parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.underline = True
        else:
            paragraph.add_run(part)
