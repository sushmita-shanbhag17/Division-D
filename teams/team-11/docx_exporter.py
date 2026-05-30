from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime, timedelta
import io

def add_colored_heading(doc, text, level=1, color=(0, 82, 155)):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_checkbox_paragraph(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run("☐  " + text)
    run.font.size = Pt(11)
    return para

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def export_to_docx(
    employee_name: str,
    current_role: str,
    target_role: str,
    gap_df: pd.DataFrame,
    learning_path: str,
    course_recommendations: dict
):
    """
    Generates a complete DOCX progress tracker.
    Returns bytes (io.BytesIO) for Streamlit download.
    """
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # =====================
    # TITLE PAGE
    # =====================
    title = doc.add_heading("", 0)
    title_run = title.add_run("Employee Skill Gap Analysis")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 82, 155)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("", 2)
    sub_run = subtitle.add_run("Personalised 6-Month Learning Path & Progress Tracker")
    sub_run.font.color.rgb = RGBColor(80, 80, 80)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Employee info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Employee Name", employee_name),
        ("Current Role", current_role),
        ("Target Role", target_role),
        ("Report Generated", datetime.now().strftime("%d %B %Y")),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.rows[i].cells[0]
        cell_value = info_table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = value
        cell_label.paragraphs[0].runs[0].font.bold = True
        shade_cell(cell_label, "D6E4F0")

    doc.add_paragraph()
    doc.add_page_break()

    # =====================
    # SECTION 1: SKILL GAP MATRIX
    # =====================
    add_colored_heading(doc, "Section 1: Skill Gap Matrix", level=1)
    doc.add_paragraph("The table below shows your current proficiency versus the requirements for your target role.")
    doc.add_paragraph()

    if not gap_df.empty:
        display_cols = ["Skill", "Category", "Current Level", "Required Level", "Gap Score", "Status"]
        available_cols = [c for c in display_cols if c in gap_df.columns]

        table = doc.add_table(rows=1 + len(gap_df), cols=len(available_cols))
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(available_cols):
            header_cells[i].text = col
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            shade_cell(header_cells[i], "00529B")
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for row_idx, (_, row) in enumerate(gap_df.iterrows()):
            cells = table.rows[row_idx + 1].cells
            for col_idx, col in enumerate(available_cols):
                cells[col_idx].text = str(row[col])

            # Color code by gap
            gap_val = row.get("Gap Score", 0)
            if gap_val == 0:
                color = "C8F7C5"  # green
            elif gap_val == 1:
                color = "FFF3CD"  # yellow
            else:
                color = "FADBD8"  # red
            for cell in cells:
                shade_cell(cell, color)
    else:
        doc.add_paragraph("No gap data available.")

    doc.add_page_break()

    # =====================
    # SECTION 2: LEARNING PATH
    # =====================
    add_colored_heading(doc, "Section 2: 6-Month Learning Path", level=1)
    doc.add_paragraph()

    if learning_path:
        for line in learning_path.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("## "):
                add_colored_heading(doc, line.replace("## ", ""), level=2)
            elif line.startswith("### "):
                add_colored_heading(doc, line.replace("### ", ""), level=3, color=(41, 128, 185))
            elif line.startswith("**Month"):
                add_colored_heading(doc, line.replace("**", ""), level=3, color=(39, 174, 96))
            elif line.startswith("- Week") or line.startswith("- Milestone"):
                add_checkbox_paragraph(doc, line[2:])
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.replace("**", ""))
                run.font.bold = True
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Learning path not generated.")

    doc.add_page_break()

    # =====================
    # SECTION 3: COURSE RECOMMENDATIONS
    # =====================
    add_colored_heading(doc, "Section 3: Course Recommendations", level=1)
    doc.add_paragraph("Recommended Udemy courses for each skill gap:")
    doc.add_paragraph()

    if course_recommendations:
        for skill, courses in course_recommendations.items():
            add_colored_heading(doc, skill, level=2, color=(142, 68, 173))
            for course in courses:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(f"📚 {course['title']}")
                run.font.bold = True
                run.font.size = Pt(11)
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.5)
                p2.add_run(f"Level: {course.get('level', 'All Levels')}  |  URL: {course.get('url', '#')}")
                p2.runs[0].font.size = Pt(10)
                p2.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No course recommendations available.")

    doc.add_page_break()

    # =====================
    # SECTION 4: PROGRESS TRACKER
    # =====================
    add_colored_heading(doc, "Section 4: Monthly Progress Tracker", level=1)
    doc.add_paragraph("Use this tracker to mark your progress each month. Check off completed milestones.")
    doc.add_paragraph()

    start_date = datetime.now()
    month_colors = ["E8F4FD", "EAF7EA", "FEF9E7", "FDEDEC", "F4ECF7", "E8F8F5"]

    for month_num in range(1, 7):
        month_date = (start_date + timedelta(days=30 * (month_num - 1))).strftime("%B %Y")
        add_colored_heading(doc, f"Month {month_num} — {month_date}", level=2, color=(41, 128, 185))

        tracker_table = doc.add_table(rows=5, cols=2)
        tracker_table.style = "Table Grid"

        headers = ["Week", "Task / Milestone"]
        for i, h in enumerate(headers):
            tracker_table.rows[0].cells[i].text = h
            tracker_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            shade_cell(tracker_table.rows[0].cells[i], "2980B9")
            tracker_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for week in range(1, 5):
            tracker_table.rows[week].cells[0].text = f"Week {week}"
            tracker_table.rows[week].cells[1].text = "☐  [Fill in your weekly task here]"
            shade_cell(tracker_table.rows[week].cells[0], month_colors[month_num - 1])

        doc.add_paragraph()

    # =====================
    # FOOTER NOTE
    # =====================
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run("Generated by D11 Employee Skill Gap Analyser | Powered by O*NET + Udemy + Grok AI")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(150, 150, 150)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer