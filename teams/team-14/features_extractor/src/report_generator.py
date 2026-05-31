"""
Groq Report Generator Engine
============================
Uses the Groq Cloud API (via LLaMA3-8B) to generate a comprehensive, executive-ready
Markdown compliance report. Includes a robust local fallback template generator.
"""

import os
import json
import httpx
from typing import Any, Dict


def generate_executive_report(device_data: Dict[str, Any]) -> str:
    """
    Generates a structured security compliance report for the IoT device.
    Uses Groq LLaMA3 completions, falling back to a local markdown template on error.
    """
    api_key = os.getenv("GROQ_API_KEY")
    
    # Extract metadata and engine details
    device_name = device_data.get("name", "Unknown Device")
    manufacturer = device_data.get("manufacturer", "Unknown Manufacturer")
    model = device_data.get("model", "N/A")
    device_type = device_data.get("type", "IoT Device")
    firmware = device_data.get("firmware", "unknown")
    
    score = device_data.get("score", 60)
    projected_score = device_data.get("projectedScore", 90)
    risk_level = device_data.get("risk", "High Risk")
    
    attack_surfaces = device_data.get("attackSurfaces") or []
    owasp_mappings = device_data.get("owaspMappings") or []
    recommendations = device_data.get("recommendations") or []
    scoring_reason = ""
    if isinstance(device_data.get("securityScoring"), dict):
        scoring_reason = device_data.get("securityScoring", {}).get("reason_for_score", "")

    # Format data summary for LLM context
    context = {
        "device_info": {
            "name": device_name,
            "manufacturer": manufacturer,
            "model": model,
            "type": device_type,
            "firmware": firmware
        },
        "score_info": {
            "current_score": score,
            "projected_score": projected_score,
            "risk_level": risk_level,
            "reason": scoring_reason
        },
        "attack_surfaces": [
            {
                "name": s.get("surface_name"),
                "category": s.get("surface_category"),
                "evidence": s.get("evidence"),
                "description": s.get("description"),
                "risk_level": s.get("risk_level")
            } for s in attack_surfaces
        ],
        "owasp_mappings": [
            {
                "category": o.get("owasp_category"),
                "severity": o.get("severity"),
                "evidence": o.get("evidence"),
                "reason": o.get("reason"),
                "affected_surfaces": o.get("affected_surfaces")
            } for o in owasp_mappings
        ],
        "recommendations": [
            {
                "title": r.get("title"),
                "severity": r.get("severity"),
                "action": r.get("action"),
                "score_reward": r.get("scoreReward")
            } for r in recommendations
        ]
    }

    if api_key:
        try:
            print("[GroqReport] Calling Groq API completions endpoint...")
            prompt = f"""You are a senior IoT cybersecurity audit reporter. 
Generate a beautifully formatted, professional executive security report in Markdown format based on the following JSON context.

REPORT TEMPLATE/FORMAT RULES:
1. Title: Executive Security & Compliance Audit: {device_name}
2. Section: Device Metadata Table (columns: Field, Detail)
3. Section: Attack Surface Exposure Analysis
   - Describe each identified exposure point/surface. Include category, risk, and exact evidence text.
4. Section: OWASP IoT Top 10 Risk Mapping
   - Match identified items with OWASP categories. Detail the reason and affected surfaces.
5. Section: Security Score & Risk Profile
   - Current Score: {score}/100 ({risk_level})
   - Include the calculated deterministic reason: {scoring_reason}
6. Section: Executive Mitigation Strategy
   - Present recommendations in order of severity (Critical/High first). Mention the projected score reward.
7. Section: Post-Mitigation Security Projection
   - Clearly state that implementing these changes elevates the score to {projected_score}/100.

CONTEXT DATA:
{json.dumps(context, indent=2)}

Return ONLY the markdown text. Keep the tone executive-ready, technical, and polished. Do not include markdown code block fences (e.g. ```markdown) in your outer response.
"""
            # Call Groq API via httpx
            with httpx.Client(timeout=30.0) as client:
                res = client.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "llama-3.1-8b-instant",
                        "messages": [
                            {"role": "user", "content": prompt}
                        ],
                        "temperature": 0.2
                    }
                )
                if res.status_code == 200:
                    report_text = res.json()["choices"][0]["message"]["content"]
                    print("[GroqReport] Report successfully generated using Groq LLaMA3.")
                    return report_text.strip()
                else:
                    print(f"[GroqReport] API returned status {res.status_code}: {res.text}")
        except Exception as e:
            print(f"[GroqReport] Error calling Groq API: {e}")

    # Fallback to local markdown generator
    print("[GroqReport] Falling back to local markdown report generator.")
    
    # Render metadata table
    meta_table = f"""
| Metadata Field | Value |
| :--- | :--- |
| **Device Name** | {device_name} |
| **Manufacturer** | {manufacturer} |
| **Model** | {model} |
| **Device Type** | {device_type} |
| **Firmware Version** | {firmware} |
"""

    # Render attack surfaces
    surfaces_md = ""
    for s in attack_surfaces:
        surfaces_md += f"""
### [Exposure] {s.get('surface_name')} ({s.get('surface_category')})
* **Risk Level:** {s.get('risk_level')}
* **Evidence:** `{s.get('evidence')}`
* **Description:** {s.get('description')}
"""

    # Render OWASP
    owasp_md = ""
    for o in owasp_mappings:
        owasp_md += f"""
### [OWASP] {o.get('owasp_category')}
* **Severity:** {o.get('severity')}
* **Evidence:** {o.get('evidence')}
* **Reason:** {o.get('reason')}
* **Affected Surfaces:** {", ".join(o.get('affected_surfaces') or []) or "None"}
"""

    # Render Recs
    recs_md = ""
    for r in recommendations:
        recs_md += f"""
* **[{r.get('severity')}] {r.get('title')}**
  * *Action:* {r.get('action')}
  * *Score Reward:* +{r.get('scoreReward')} points
"""

    local_report = f"""# Executive Security & Compliance Audit: {device_name}

## [Info] Device Information
{meta_table}

---

## [Exposure] Attack Surface Exposure Analysis
{surfaces_md or "*No attack surfaces exposed.*"}

---

## [OWASP] OWASP IoT Top 10 Risk Mapping
{owasp_md or "*No OWASP IoT Top 10 risks mapped.*"}

---

## [Score] Security Score & Risk Profile
* **Security Score:** **{score}/100**
* **Assessed Risk Level:** **{risk_level}**

### **Reason for Score**
{scoring_reason or f"Assigned based on detected insecure protocols and interfaces, with {len(attack_surfaces)} exposure points and {len(owasp_mappings)} OWASP risk flags."}

---

## [Mitigation] Executive Mitigation Strategy
Here are the recommended remediations to secure the device:
{recs_md or "*No mitigations required.*"}

---

## [Projection] Post-Mitigation Security Projection
By executing all the recommended mitigation policies, the device security score is projected to improve to:
* **Projected Security Score:** **{projected_score}/100 (Low Risk)**
"""
    return local_report


import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF

class PDFReport(FPDF):
    def header(self):
        pass
        
    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

def export_report_to_pdf(markdown_text: str) -> bytes:
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=10)
    
    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Heading 1
        if line.startswith('# '):
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(0, 120, 180) # Sleek blue
            pdf.multi_cell(0, 10, line[2:])
            pdf.ln(4)
        # Heading 2
        elif line.startswith('## '):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 8, line[3:])
            pdf.ln(2)
        # Heading 3
        elif line.startswith('### '):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(80, 80, 80)
            pdf.multi_cell(0, 6, line[4:])
            pdf.ln(1)
        # Horizontal rule
        elif line.startswith('---'):
            pdf.ln(2)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(4)
        # Table
        elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|') and '---' in lines[i + 1]:
            headers = [col.strip() for col in line.split('|')[1:-1]]
            
            # Read rows
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].startswith('|'):
                row_cells = [col.strip() for col in lines[j].split('|')[1:-1]]
                rows.append(row_cells)
                j += 1
                
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(240, 240, 240)
            pdf.set_text_color(0, 0, 0)
            pdf.set_draw_color(220, 220, 220)
            
            col_width = 190.0 / max(len(headers), 1)
            
            # Draw header row
            for h in headers:
                pdf.cell(col_width, 8, h, border=1, fill=True, align="L")
            pdf.ln(8)
            
            pdf.set_font("Helvetica", size=8)
            for row in rows:
                for idx, cell_val in enumerate(row):
                    clean_cell = cell_val.replace('**', '').replace('`', '')
                    if idx < len(headers):
                        pdf.cell(col_width, 7, clean_cell[:40], border=1, align="L")
                pdf.ln(7)
                
            pdf.ln(4)
            i = j - 1
        # Bullet list
        elif line.startswith('* ') or line.startswith('- '):
            pdf.set_font("Helvetica", size=9)
            pdf.set_text_color(60, 60, 60)
            clean_text = line[2:].replace('**', '').replace('`', '')
            pdf.multi_cell(0, 6, f"- {clean_text}")
            pdf.ln(1)
        # Indented bullet
        elif line.startswith('  * ') or line.startswith('  - '):
            pdf.set_font("Helvetica", size=9)
            pdf.set_text_color(80, 80, 80)
            clean_text = line[4:].replace('**', '').replace('`', '')
            pdf.multi_cell(0, 6, f"  * {clean_text}")
            pdf.ln(1)
        # Paragraph
        elif line.strip():
            pdf.set_font("Helvetica", size=9.5)
            pdf.set_text_color(40, 40, 40)
            clean_text = line.replace('**', '').replace('`', '')
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(2)
            
        i += 1
        
    return bytes(pdf.output())

def export_report_to_docx(markdown_text: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Heading 1
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(6)
        # Heading 2
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(4)
        # Heading 3
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(2)
        # Horizontal rule
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p_run = p.add_run('━' * 40)
            p_run.font.color.rgb = RGBColor(128, 128, 128)
            p.paragraph_format.space_after = Pt(8)
        # Table
        elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|') and '---' in lines[i + 1]:
            headers = [col.strip() for col in line.split('|')[1:-1]]
            
            # Read rows
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].startswith('|'):
                row_cells = [col.strip() for col in lines[j].split('|')[1:-1]]
                rows.append(row_cells)
                j += 1
                
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Shading Accent 1'
            
            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                hdr_cells[idx].text = header
                for p in hdr_cells[idx].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        
            for row_data in rows:
                row_cells = table.add_row().cells
                for idx, cell_val in enumerate(row_data):
                    if idx < len(row_cells):
                        row_cells[idx].text = cell_val
                        
            doc.add_paragraph()
            i = j - 1
        # Bullet list
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)
        # Indented bullet
        elif line.startswith('  * ') or line.startswith('  - '):
            text = line[4:]
            p = doc.add_paragraph(style='List Bullet 2')
            _add_formatted_runs(p, text)
        # Paragraph
        elif line.strip():
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
            
        i += 1
        
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

def _add_formatted_runs(paragraph, text: str):
    import re
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)
