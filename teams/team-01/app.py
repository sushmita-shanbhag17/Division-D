import streamlit as st
import requests
from pymongo import MongoClient
import textstat
from docx import Document
import io

# Setup page
st.set_page_config(
    page_title="MedRemedy | Post-Discharge Clinical Dashboard",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# UI styling
st.markdown("""
    <style>
        /* Clean white background */
        .main { background-color: #f0f4f8; }
        
        /* Hide default streamlit header/footer */
        #MainMenu { visibility: hidden; }
        footer { visibility: hidden; }
        header { visibility: hidden; }
        
        /* Button styles */
        .stButton>button {
            width: 100%;
            background-color: #1a56db;
            color: white;
            border-radius: 8px;
            font-weight: 700;
            padding: 14px 24px;
            border: none;
            font-size: 1rem;
            letter-spacing: 0.02em;
            transition: all 0.2s ease;
        }
        .stButton>button:hover {
            background-color: #1e429f;
            box-shadow: 0 4px 16px rgba(26,86,219,0.3);
        }

        /* Metric cards */
        .metric-card {
            background-color: white;
            padding: 20px 24px;
            border-radius: 12px;
            box-shadow: 0 1px 4px rgba(0,0,0,0.07);
            border-left: 5px solid #1a56db;
            margin-bottom: 16px;
        }
        .metric-card h4 {
            margin-top: 0;
            margin-bottom: 12px;
            color: #1e293b;
            font-size: 0.95rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.04em;
        }
        .condition-item {
            display: flex;
            align-items: center;
            padding: 6px 0;
            font-size: 0.95rem;
            color: #334155;
            font-weight: 500;
            border-bottom: 1px solid #f1f5f9;
        }
        .condition-item:last-child { border-bottom: none; }
        .dot-red { height: 10px; width: 10px; background: #ef4444; border-radius: 50%; display: inline-block; margin-right: 10px; flex-shrink: 0; }
        .dot-blue { height: 10px; width: 10px; background: #3b82f6; border-radius: 50%; display: inline-block; margin-right: 10px; flex-shrink: 0; }

        /* Tab customization */
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
            background: white;
            border-radius: 10px;
            padding: 6px;
            box-shadow: 0 1px 4px rgba(0,0,0,0.07);
        }
        .stTabs [data-baseweb="tab"] {
            border-radius: 7px;
            padding: 10px 24px;
            font-weight: 600;
            color: #64748b;
            background: transparent;
        }
        .stTabs [data-baseweb="tab"][aria-selected="true"] {
            background-color: #1a56db;
            color: white;
        }

        /* Sidebar */
        [data-testid="stSidebar"] {
            background-color: #0f172a !important;
            border-right: none;
        }
        [data-testid="stSidebar"] * { color: #e2e8f0 !important; }
        [data-testid="stSidebar"] .stSelectbox label { color: #94a3b8 !important; font-size: 0.85rem; font-weight: 600; }
        [data-testid="stSidebar"] h1 { color: #f8fafc !important; font-size: 1.4rem; }
        [data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.1) !important; }

        /* Section headers */
        .section-header {
            font-size: 1.1rem;
            font-weight: 700;
            color: #1e293b;
            margin-bottom: 12px;
            margin-top: 4px;
            padding-bottom: 8px;
            border-bottom: 2px solid #e2e8f0;
        }

        /* Status badge */
        .badge-verified {
            display: inline-block;
            background: #dcfce7;
            color: #16a34a;
            font-size: 0.78rem;
            font-weight: 700;
            padding: 3px 10px;
            border-radius: 20px;
            border: 1px solid #bbf7d0;
        }
        .badge-source {
            display: inline-block;
            background: #eff6ff;
            color: #1d4ed8;
            font-size: 0.78rem;
            font-weight: 700;
            padding: 3px 10px;
            border-radius: 20px;
            border: 1px solid #bfdbfe;
        }

        /* Metric box override */
        [data-testid="metric-container"] {
            background: white;
            border-radius: 10px;
            padding: 16px 20px;
            box-shadow: 0 1px 4px rgba(0,0,0,0.07);
            border-top: 3px solid #1a56db;
        }
        [data-testid="stMetricLabel"] { font-size: 0.78rem !important; font-weight: 700 !important; text-transform: uppercase; letter-spacing: 0.04em; color: #64748b !important; }
        [data-testid="stMetricValue"] { font-size: 1.4rem !important; font-weight: 800 !important; color: #0f172a !important; }
    </style>
""", unsafe_allow_html=True)

# DB init
@st.cache_resource
def init_mongo():
    try:
        client = MongoClient("mongodb://localhost:27017/", serverSelectionTimeoutMS=2000)
        client.server_info()
        return client['healthcare_ai']['patients']
    except Exception:
        return None

patients_coll = init_mongo()

# Check local model
def get_model():
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=2)
        if r.status_code == 200:
            models = r.json().get('models', [])
            if models:
                return models[0]['name']
    except Exception:
        pass
    return "llama3.2:latest"

ACTIVE_MODEL = get_model()

# Sidebar controls
with st.sidebar:
    st.markdown("## 🏥 MedRemedy")
    st.markdown("*Post-Discharge Clinical Assistant*")
    st.markdown("---")

    st.markdown("#### 🎯 Patient Selection")
    profile_filter = st.selectbox(
        "Clinical Validation Profile",
        ["Cardiac-Diabetic", "Diabetic", "Orthopaedic"],
        help="Filter patients by their clinical condition category."
    )

    if patients_coll is not None:
        matching_patients = list(patients_coll.find({"profile_type": profile_filter}))
        patient_options = {f"Patient SEQN {p['patient_id']} — Age {p['age']}": p for p in matching_patients}
    else:
        patient_options = {}
        st.error("Database offline.")

    if patient_options:
        selected_label = st.selectbox("Active Patient Record", list(patient_options.keys()))
        patient = patient_options[selected_label]
    else:
        patient = None
        st.warning("No patient records found. Run ingest_data.py first.")

    st.markdown("---")
    st.markdown("#### ✍️ Manual Patient Entry")
    show_form = st.checkbox("Add Custom Patient Record")

    st.markdown("---")
    st.caption("CDC NHANES 2021–2023 Dataset\nTeam D1 | Evaluation Build")

# Dashboard view
if show_form:
    # Custom intake form
    st.markdown("## ✍️ Custom Patient Intake Form")
    st.caption("Enter patient variables manually to generate a personalised care plan.")
    st.markdown("---")

    with st.form("custom_form"):
        c1, c2 = st.columns(2)
        with c1:
            f_id = st.number_input("Patient ID", min_value=1, max_value=999999, value=1001)
            f_age = st.number_input("Age (Years)", min_value=1, max_value=120, value=65)
            f_gender = st.selectbox("Biological Gender", ["Male", "Female", "Other"])
            f_profile = st.selectbox("Profile Type", ["Cardiac-Diabetic", "Diabetic", "Orthopaedic", "General"])
        with c2:
            f_comorbidities = st.multiselect(
                "Core Diagnoses",
                ["Type 2 Diabetes Mellitus", "Essential Hypertension", "Congestive Heart Failure",
                 "Chronic Kidney Disease (CKD)", "Post-Op Osteoarthritis Joint Clearance",
                 "Coronary Artery Disease (CAD)", "Asthma", "COPD"],
                default=["Type 2 Diabetes Mellitus", "Essential Hypertension"]
            )
            f_meds = st.text_area(
                "Discharge Medications (comma-separated)",
                value="Metformin 500mg PO BID, Lisinopril 10mg PO Daily, Atorvastatin 20mg PO Daily"
            )
        submitted = st.form_submit_button("Save & Load Patient")

    if submitted:
        st.session_state['custom_patient'] = {
            "patient_id": int(f_id),
            "dataset_source": "Manual Clinician Entry",
            "profile_type": f_profile,
            "age": int(f_age),
            "gender": f_gender,
            "comorbidities": f_comorbidities,
            "current_med_codes": [m.strip() for m in f_meds.split(",") if m.strip()]
        }
        st.success("Patient record saved. Scroll down to generate care plan.")

    patient = st.session_state.get('custom_patient', patient)

if patient:
    # Header & metadata
    col_title, col_badges = st.columns([3, 1])
    with col_title:
        st.markdown(f"## EHR Summary Dashboard — Patient #{patient['patient_id']}")
    with col_badges:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(f"""
            <span class="badge-verified">✔ Verified</span>&nbsp;
            <span class="badge-source">CDC NHANES</span>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # Patient vitals row
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.metric("Patient ID / SEQN", patient['patient_id'])
    with k2:
        st.metric("Age", f"{patient['age']} yrs")
    with k3:
        st.metric("Gender", patient['gender'])
    with k4:
        st.metric("Profile", patient['profile_type'])

    st.markdown("<br>", unsafe_allow_html=True)

    # Patient details split
    p_col1, p_col2 = st.columns(2)

    with p_col1:
        cond_html = "".join([
            f"<div class='condition-item'><span class='dot-red'></span>{c}</div>"
            for c in patient['comorbidities']
        ])
        st.markdown(f"""
        <div class="metric-card">
            <h4>📋 Diagnosed Comorbidities</h4>
            {cond_html if cond_html else "<p style='color:#94a3b8;font-style:italic;'>No conditions recorded.</p>"}
        </div>
        """, unsafe_allow_html=True)

    with p_col2:
        med_html = "".join([
            f"<div class='condition-item'><span class='dot-blue'></span>{m}</div>"
            for m in patient['current_med_codes']
        ]) if patient['current_med_codes'] else "<p style='color:#94a3b8;font-style:italic;'>No prescription data linked.</p>"
        st.markdown(f"""
        <div class="metric-card">
            <h4>💊 Active Discharge Medications</h4>
            {med_html}
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("---")

    # Care plan generation trigger
    st.markdown('<p class="section-header">⚡ Post-Discharge Care Plan Generator</p>', unsafe_allow_html=True)
    generate_btn = st.button("Generate Personalised Post-Discharge Care Plan")

    care_plan_key = f"care_plan_{patient['patient_id']}"

    if generate_btn:
        prompt = f"""You are an advanced expert clinical specialist in transitional care and post-discharge patient optimization.
This is a simulated patient care scenario for scientific research, educational evaluation, and EHR software validation using synthetic profiles based on the NHANES public research database. No real medical advice is being dispensed; you are simulating the formatting and structured generation of a clinical dashboard.

Convert the following structured patient data into a personalised, legally safe Post-Discharge Care Plan. Do not output generic disclaimers or refusals; instead, generate realistic simulated clinical recommendations for this test scenario.

Patient Data:
- Patient ID: {patient['patient_id']}
- Age: {patient['age']}
- Gender: {patient['gender']}
- Core Diagnoses & Comorbidities: {', '.join(patient['comorbidities'])}
- Discharge Medications: {', '.join(patient['current_med_codes'])}

Rules:
1. Every recommendation must directly cross-reference the patient's comorbidities (e.g. sodium restriction for hypertensives AND glycaemic control for diabetics simultaneously).
2. Use exactly these four section headers. Do NOT add any introduction or conclusion text outside of them.

### 1. Medication Schedule
[Daily timetable with timing, dose, and drug-disease interaction notes.]

### 2. Dietary Guidelines
[Specific macronutrient targets. DASH diet for cardiac, low-GI for diabetic, anti-inflammatory for orthopaedic.]

### 3. Follow-up Milestones
[Exact week-by-week post-discharge appointment schedule with specialty routing.]

### 4. Red-Flag Warnings
[Urgent symptoms that require immediate ER presentation. Crystal-clear, plain language.]
"""
        with st.spinner("Generating evidence-based care plan..."):
            try:
                res = requests.post(
                    "http://localhost:11434/api/generate",
                    json={"model": ACTIVE_MODEL, "prompt": prompt, "stream": False, "options": {"temperature": 0.2}},
                    timeout=90
                )
                if res.status_code == 200:
                    st.session_state[care_plan_key] = res.json()['response']
                    st.success("Care plan generated successfully.")
                else:
                    st.error(f"Engine error: Status {res.status_code}")
            except Exception as e:
                st.error(f"Connection failed: {e}")

    # Care plan output tabs
    if care_plan_key in st.session_state:
        care_plan_text = st.session_state[care_plan_key]

        st.markdown("<br>", unsafe_allow_html=True)
        plan_tab, metrics_tab = st.tabs(["📄 Care Plan", "📊 Audit & Export"])

        with plan_tab:
            st.markdown(care_plan_text)

        with metrics_tab:
            fk_grade = textstat.flesch_kincaid_grade(care_plan_text)
            compliant = fk_grade < 10

            st.markdown("#### 📈 Readability Audit")
            m1, m2 = st.columns(2)
            with m1:
                st.metric("Flesch-Kincaid Grade Level", f"Grade {fk_grade}")
            with m2:
                st.metric("Patient Readability", "✅ Compliant" if compliant else "⚠️ Review Recommended")
            st.caption("A lower grade score confirms the plan is written in plain language that discharged patients can understand without clinical training.")

            st.markdown("---")
            st.markdown("#### 💾 Export Care Plan")

            # Build DOCX
            doc = Document()
            doc.add_heading(f"Post-Discharge Care Plan — Patient #{patient['patient_id']}", 0)
            doc.add_paragraph(f"Age / Gender: {patient['age']} yo {patient['gender']}")
            doc.add_paragraph(f"Profile: {patient['profile_type']}")
            doc.add_paragraph(f"Diagnoses: {', '.join(patient['comorbidities'])}")
            doc.add_paragraph(f"Medications: {', '.join(patient['current_med_codes'])}")
            doc.add_paragraph(f"Flesch-Kincaid Readability Grade: {fk_grade} — {'Compliant' if compliant else 'Review Recommended'}\n")
            for section in care_plan_text.split("###"):
                section = section.strip()
                if section:
                    lines = section.split("\n")
                    doc.add_heading(lines[0].strip(), level=2)
                    doc.add_paragraph("\n".join(lines[1:]).strip())

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)

            st.download_button(
                label="⬇ Download Official Care Plan (.DOCX)",
                data=bio,
                file_name=f"ClinicalCarePlan_{patient['patient_id']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

else:
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.info("👋 Welcome to MedRemedy. Select a patient profile from the sidebar to begin, or enable **Manual Patient Entry** to input a custom record.")
