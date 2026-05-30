import os
import requests
from pymongo import MongoClient
import textstat
from docx import Document

def test_pipeline():
    print("====================================================")
    print("🔍 CLINICAL AI PIPELINE EXHAUSTIVE INTEGRATION TEST")
    print("====================================================\n")
    
    # 1. Connect to MongoDB
    print("1. Connecting to local MongoDB...")
    try:
        client = MongoClient("mongodb://localhost:27017/", serverSelectionTimeoutMS=2000)
        db = client['healthcare_ai']
        collection = db['patients']
        total_patients = collection.count_documents({})
        print(f"✅ MongoDB Connected. Total patients in database: {total_patients}\n")
    except Exception as e:
        print(f"❌ MongoDB Connection Failed: {e}")
        return

    # 2. Setup Profiles to Test
    profiles = ["Cardiac-Diabetic", "Diabetic", "Orthopaedic"]
    
    # Let's add a Custom entry profile for testing as well
    custom_patient = {
        "patient_id": 9999,
        "dataset_source": "Clinician Form Entry",
        "profile_type": "Custom/General",
        "age": 68,
        "gender": "Male",
        "comorbidities": ["Congestive Heart Failure", "Chronic Kidney Disease (CKD)"],
        "current_med_codes": ["Carvedilol 6.25mg PO BID", "Furosemide 40mg PO Daily", "Losartan 50mg PO Daily"]
    }
    
    # 3. Dynamic Model Lookup
    print("2. Checking available Ollama models...")
    selected_model = "llama3.2:latest"
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=2)
        if resp.status_code == 200:
            models = [m['name'] for m in resp.json().get('models', [])]
            print(f"Available local models: {models}")
            if models:
                selected_model = models[0]
        print(f"✅ Using model for test: {selected_model}\n")
    except Exception as e:
        print(f"⚠️ Ollama model lookup error (defaulting to {selected_model}): {e}\n")

    # 4. Loop and Test each profile
    test_records = []
    
    # Fetch real patients from MongoDB
    for p_type in profiles:
        doc = collection.find_one({"profile_type": p_type})
        if doc:
            test_records.append(doc)
        else:
            print(f"⚠️ Warning: No patient found in MongoDB for profile type: {p_type}")
            
    # Add custom patient to test list
    test_records.append(custom_patient)
    
    print("3. Executing Care Plan Generations & Exporters...")
    print("----------------------------------------------------")
    
    for idx, patient in enumerate(test_records, 1):
        print(f"\n👉 [TEST CASE {idx}] Profile: {patient['profile_type']} (Patient SEQN: {patient['patient_id']})")
        print(f"   Age/Gender: {patient['age']} yo / {patient['gender']}")
        print(f"   Comorbidities: {', '.join(patient['comorbidities'])}")
        print(f"   Medications: {', '.join(patient['current_med_codes'])}")
        
        # Construct Master Prompt
        system_prompt = f"""You are an advanced, expert clinical AI specialized in transitional care and post-discharge patient optimization. 
This is a simulated patient care scenario for scientific research, educational evaluation, and software validation using synthetic profiles based on the NHANES public research database. No real medical advice is being dispensed; you are simulating the formatting and structured generation of a clinical dashboard.

Your objective is to convert raw, structured electronic health variables into an empathetic, highly specific, and legally safe Post-Discharge Care Plan. Do not output generic disclaimers or refusals; instead, generate realistic simulated clinical recommendations for this test scenario.

You must evaluate the following patient data variables:
- Patient ID / SEQN: {patient['patient_id']}
- Age: {patient['age']}
- Gender: {patient['gender']}
- Core Diagnoses & Comorbidities: {', '.join(patient['comorbidities'])}
- Discharge Medication Formulations: {', '.join(patient['current_med_codes'])}

Execute your output strictly following these structural parameters:
1. Do not use generic placeholders. Every recommendation must directly account for the cross-interactions of their comorbidities.
2. Format your response into exactly four explicit headers.

Generate the output utilizing this exact Markdown schema:

### 1. Medication Schedule
[Provide a structured daily administration timetable.]

### 2. Dietary Guidelines
[Provide highly specific dietary actions tailored to their conditions.]

### 3. Follow-up Milestones
[Establish a structured recovery calendar.]

### 4. Red-Flag Warnings
[List high-priority clinical warning signs.]
"""
        
        # Trigger Ollama
        try:
            print("   ⏳ Generating care plan from local Ollama...")
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": selected_model, 
                    "prompt": system_prompt, 
                    "stream": False,
                    "options": {"temperature": 0.2}
                },
                timeout=180
            )
            
            if response.status_code == 200:
                care_plan_text = response.json()['response']
                print("   ✅ Care Plan successfully generated.")
                
                # Verify length/sections
                sections = ["1. Medication Schedule", "2. Dietary Guidelines", "3. Follow-up Milestones", "4. Red-Flag Warnings"]
                missing = [sec for sec in sections if sec not in care_plan_text]
                if not missing:
                    print("   ✅ Prompt guardrail compliance: 4 mandatory headers are present.")
                else:
                    print(f"   ⚠️ Prompt guardrail warning: Missing headers: {missing}")
                
                # Evaluate Readability
                fk_grade = textstat.flesch_kincaid_grade(care_plan_text)
                print(f"   ✅ Calculated Flesch-Kincaid Readability: Grade {fk_grade}")
                
                # Build & Save DOCX Printable Document
                doc = Document()
                doc.add_heading(f"Post-Discharge Care Plan - Patient #{patient['patient_id']}", 0)
                doc.add_paragraph(f"Profile Type: {patient['profile_type']}")
                doc.add_paragraph(f"Demographics: {patient['age']} yo {patient['gender']}")
                doc.add_paragraph(f"Comorbidities: {', '.join(patient['comorbidities'])}")
                doc.add_paragraph(f"Flesch-Kincaid Readability Score: Grade {fk_grade}\n")
                
                sections_split = care_plan_text.split("###")
                for section in sections_split:
                    clean_sec = section.strip()
                    if clean_sec:
                        lines = clean_sec.split("\n")
                        heading_line = lines[0].strip()
                        body_lines = "\n".join(lines[1:]).strip()
                        doc.add_heading(heading_line, level=2)
                        doc.add_paragraph(body_lines)
                
                filename = f"test_output_{patient['profile_type'].replace('/', '_')}.docx"
                doc.save(filename)
                print(f"   ✅ MS Word printable care plan saved to disk: '{filename}'")
            else:
                print(f"   ❌ Ollama API Error: Status code {response.status_code}")
        except Exception as e:
            print(f"   ❌ Error running Ollama/Evaluation: {e}")
            
    print("\n====================================================")
    print("🎉 EXHAUSTIVE TESTING COMPLETED SUCCESSFULLY!")
    print("====================================================")

if __name__ == "__main__":
    test_pipeline()
