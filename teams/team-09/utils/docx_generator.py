from docx import Document


def save_docx(content, filename):

    doc = Document()

    doc.add_heading(
        "Cold Chain Breach Alert Report",
        level=1
    )

    doc.add_paragraph(content)

    doc.save(filename)

    print(f"{filename} generated successfully")